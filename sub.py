from docx import Document
from docx.shared import Inches
import datetime
import os
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from docx.shared import Pt
from flask import *
from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
import smtplib

app = Flask(__name__)

@app.route('/home',methods = ['GET','POST'])
def home():
    global prtype
    global prno
    if request.method =='POST':
        file = request.files['file[]']
        prtype=request.form['type']
        prno=request.form['prlist']
        print(prtype + " " + prno)
        if file:
            filename = "ab.docx"
            file.save(os.path.join('./uploads',filename))
            gsheet()
    prlist=[1,2,3,4,5,6,7,8,9,10,11,12]
    return render_template("home.html",prlist=prlist)
            
def createop(student):
    document = Document('./uploads/ab.docx')
    date = datetime.datetime.now()
    date=date.strftime("%d/%m/%Y")
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[1]
    line0=header.paragraphs[0].text
    line=str(line0).split(" ")
    obj = filter(lambda x: x != "", line)
    line=list(obj)
    header.paragraphs[0].text=line[0]+" "+line[1]+"\t  "+line[2]+" "+line[3]+" "+line[4]
    paragraph.text = "Enrollment No: "+ student[0]+"\t  Name: "+student[1]+"\t  Date: "+ date
    '''style = document.styles['Normal']
    font = style.font
    font.name ='Arial'
    font.size = Pt(11)
    print(paragraph.text)'''
    filename=student[0]+"_"+prtype+prno
    document.save(filename+".docx")
    #os.system("abiword --to=pdf demo.odt")

    os.system("libreoffice --headless --convert-to pdf --outdir . "+ filename + ".docx")
    #mail(student[2],filename+".pdf")


def gsheet():
    scope = ['https://spreadsheets.google.com/feeds',
    'https://www.googleapis.com/auth/drive','https://www.googleapis.com/auth/spreadsheets',]
    creds = ServiceAccountCredentials.from_json_keyfile_name('Result Analysis-a62d666095bf.json', scope)
    client = gspread.authorize(creds)
    sheet = client.open("Assignment").sheet1
    print(sheet)
    #s = sheet.get("{SPREADSHEET_URL}")
    #print(s) # will ensure your file is accessible 
    #s.sheet[1].to_csv('Spam.csv', encoding='utf-8', dialect='excel')
    stu_list=sheet.get_all_values()
    #print(stu_list)
    stu_list.pop(0)
    for i in stu_list:
        createop(i)

def mail(y,filename):
    sender_address = 'resultnotifiergtu@gmail.com'
    sender_pass = 'sarvajanikgtu'
    receiver_address=y
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] =  "testing phase (please ignore)"
    attach_file_name = filename
    attach_file = open(attach_file_name, 'rb')
    payload = MIMEBase('application', 'octate-stream')
    payload.set_payload((attach_file).read())
    encoders.encode_base64(payload)
    payload.add_header('Content-Disposition', 'attachment', filename=attach_file_name)
    message.attach(payload)
    session = smtplib.SMTP('smtp.gmail.com', 587) #use gmail with port
    session.starttls() #enable security
    session.login(sender_address, sender_pass) #login with mail_id and password
    text = message.as_string()
    session.sendmail(sender_address, receiver_address, text)
    session.quit()
    print('Mail Sent')

if __name__ == '__main__':
    app.run(debug = True)
